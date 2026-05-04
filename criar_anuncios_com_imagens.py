#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import pdfplumber
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from pathlib import Path

# Caminhos
pdf_path = "/sessions/gallant-magical-clarke/mnt/uploads/MDL TABLOIDE MENSAL  (32).pdf"
output_dir = "/sessions/gallant-magical-clarke/mnt/outputs"
images_dir = os.path.join(output_dir, "imagens_temp")

# Criar diretório temporário para imagens
os.makedirs(images_dir, exist_ok=True)

print("Extrai imagens do PDF...")

try:
    # Converter PDF para imagens
    images = convert_from_path(pdf_path, dpi=100)

    imagem_paths = []
    for i, image in enumerate(images, 1):
        image_path = os.path.join(images_dir, f"pagina_{i}.jpg")
        image.save(image_path, 'JPEG', quality=85)
        imagem_paths.append(image_path)
        print(f"✓ Página {i} extraída")

    # Criar documento Word
    doc = Document()

    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Adicionar título
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("🎯 ANÚNCIOS COM IMAGENS DOS PRODUTOS\nMDL - Móveis do Lar")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(255, 102, 0)

    # Adicionar as páginas do PDF como imagens
    for idx, img_path in enumerate(imagem_paths, 1):
        doc.add_paragraph()

        # Adicionar cabeçalho da página
        header = doc.add_paragraph()
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header.add_run(f"PÁGINA {idx} DO TABLOIDE ORIGINAL")
        header_run.font.size = Pt(12)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(255, 102, 0)

        # Adicionar imagem
        try:
            doc.add_picture(img_path, width=Inches(6.5))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            print(f"✓ Imagem da página {idx} adicionada")
        except Exception as e:
            print(f"⚠️ Erro ao adicionar imagem da página {idx}: {e}")

        # Adicionar quebra de página
        if idx < len(imagem_paths):
            doc.add_page_break()

    # Salvar documento
    doc.save(os.path.join(output_dir, 'Tabloide_com_Imagens.docx'))
    print(f"\n✓ Documento criado: Tabloide_com_Imagens.docx")

except ImportError as e:
    print(f"Erro de importação: {e}")
    print("Instale: pip install pdf2image Pillow python-docx --break-system-packages")
except Exception as e:
    print(f"Erro: {e}")

# Limpar diretório temporário
import shutil
if os.path.exists(images_dir):
    shutil.rmtree(images_dir)
    print("✓ Limpeza temporária concluída")
