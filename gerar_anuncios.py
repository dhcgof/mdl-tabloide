#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Script para extrair imagens do PDF e criar documento Word com anúncios
MDL - Móveis do Lar
"""

import os
import sys
import shutil
from pathlib import Path

# Tentar importar bibliotecas necessárias
try:
    from pdf2image import convert_from_path
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError as e:
    print(f"❌ Erro de importação: {e}")
    print("\n📦 Instalando bibliotecas necessárias...")
    os.system("pip install pdf2image Pillow python-docx pdfplumber --break-system-packages")
    print("\n⚠️  Por favor, execute o script novamente após a instalação")
    sys.exit(1)

def encontrar_pdf():
    """Encontra o arquivo PDF na pasta atual"""
    for arquivo in Path('.').glob('*.pdf'):
        return str(arquivo)
    return None

def extrair_e_criar_doc(pdf_path, dpi=150):
    """Extrai imagens do PDF e cria documento Word com anúncios"""

    print(f"\n🎯 Processando: {pdf_path}")
    print("=" * 60)

    # Diretório temporário para imagens
    temp_dir = Path('temp_images')
    if temp_dir.exists():
        shutil.rmtree(temp_dir)
    temp_dir.mkdir(exist_ok=True)

    try:
        print(f"📄 Convertendo PDF para imagens (DPI: {dpi})...")
        images = convert_from_path(pdf_path, dpi=dpi)
        print(f"✓ {len(images)} página(s) encontrada(s)")

        # Salvar imagens temporariamente
        image_paths = []
        for i, image in enumerate(images, 1):
            image_path = temp_dir / f"pagina_{i}.jpg"
            image.save(image_path, 'JPEG', quality=95)
            image_paths.append(str(image_path))
            print(f"  ✓ Página {i} extraída")

        # Criar documento Word
        print("\n📝 Criando documento Word...")
        doc = Document()

        # Configurar margens
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Título
        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.add_run("🎯 ANÚNCIOS COM IMAGENS DOS PRODUTOS\nMDL - Móveis do Lar")
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(255, 102, 0)

        # Data
        from datetime import datetime
        data_para = doc.add_paragraph()
        data_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        data_run = data_para.add_run(f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
        data_run.font.size = Pt(10)
        data_run.font.italic = True

        doc.add_paragraph()  # Espaço

        # Adicionar imagens e cabeçalhos
        for idx, img_path in enumerate(image_paths, 1):
            # Cabeçalho da página
            header = doc.add_paragraph()
            header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_run = header.add_run(f"PÁGINA {idx} - TABELA DE PRODUTOS")
            header_run.font.size = Pt(12)
            header_run.font.bold = True
            header_run.font.color.rgb = RGBColor(255, 102, 0)

            # Imagem
            try:
                doc.add_picture(img_path, width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                print(f"  ✓ Imagem página {idx} adicionada")

                # Quebra de página
                if idx < len(image_paths):
                    doc.add_page_break()

            except Exception as e:
                print(f"  ⚠️  Erro ao adicionar imagem {idx}: {e}")

        # Salvar documento
        output_file = 'Anuncios_com_Imagens.docx'
        doc.save(output_file)

        print(f"\n✅ SUCESSO!")
        print(f"📄 Arquivo criado: {output_file}")
        print(f"📍 Localização: {Path(output_file).absolute()}")
        print("=" * 60)

        return True

    except Exception as e:
        print(f"\n❌ Erro durante processamento: {e}")
        return False

    finally:
        # Limpar pasta temporária
        if temp_dir.exists():
            shutil.rmtree(temp_dir)
            print("🧹 Limpeza temporária concluída")

def main():
    """Função principal"""
    print("\n🎯 GERADOR DE ANÚNCIOS - MDL Móveis do Lar")
    print("=" * 60)

    # Encontrar PDF
    pdf_file = encontrar_pdf()

    if not pdf_file:
        print("❌ Nenhum arquivo PDF encontrado na pasta atual!")
        print("💡 Dica: Coloque o arquivo PDF na mesma pasta do script")
        return False

    # Processar
    sucesso = extrair_e_criar_doc(pdf_file, dpi=150)

    if sucesso:
        print("\n💡 Próximos passos:")
        print("   1. Abra o arquivo .docx no Microsoft Word")
        print("   2. Verifique as imagens e edite conforme necessário")
        print("   3. Exporte para PDF ou imprima")

    return sucesso

if __name__ == "__main__":
    sucesso = main()
    sys.exit(0 if sucesso else 1)
