#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Criar documento
doc = Document()

# Configurar margens
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Adicionar título
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_run = title.add_run("🎯 ANÚNCIOS INDIVIDUAIS DE PRODUTOS\nMDL - Móveis do Lar")
title_run.font.size = Pt(16)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(255, 102, 0)  # Laranja

# Adicionar data
subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_run = subtitle.add_run("Maio de 2026 • Ofertas Imperdíveis!")
subtitle_run.font.size = Pt(10)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()  # Espaço

# Lista de produtos com dados
produtos = [
    {
        "nome": "Refrigerador 2 Portas Cycle Defrost 294L",
        "marca": "MIDEA",
        "modelo": "MDRT411FGF011",
        "especificacoes": "Branco, 127V, Cycle Defrost",
        "preco": 2199.00,
        "parcelamento": "1+14x",
        "descricao": "Mantenha seus alimentos frescos por mais tempo com tecnologia Cycle Defrost. Design moderno e eficiente para sua cozinha.",
        "destaque": "PROMOÇÃO ESPECIAL"
    },
    {
        "nome": "Auto Rádio 10\" Central Multimídia",
        "marca": "ROADSTAR",
        "modelo": "RS10200BR",
        "especificacoes": "Tela 10 polegadas, Bluetooth, USB",
        "preco": 549.00,
        "parcelamento": "1+14x",
        "descricao": "Transforme sua experiência no carro com essa central multimídia moderna. Som de qualidade e conectividade total.",
        "destaque": "TECNOLOGIA"
    },
    {
        "nome": "Ferro a Vapor Bravo Cerâmico",
        "marca": "AGRATTO",
        "modelo": "FC-01",
        "especificacoes": "127V, Placa Cerâmica Antaderente",
        "preco": 69.90,
        "parcelamento": "À vista",
        "descricao": "Roupas sempre perfeitas! Com placa cerâmica que desliza suavemente, deixando suas roupas sem rugas.",
        "destaque": "SUPER PROMOÇÃO"
    },
    {
        "nome": "Lavadora 10kg Semi-Automática",
        "marca": "LIBELL",
        "modelo": "12661/12662",
        "especificacoes": "10kg, Branca, 127V",
        "preco": 479.00,
        "parcelamento": "1+14x",
        "descricao": "Lavagem potente e eficiente. Ideal para famílias que buscam praticidade e economia de água e energia.",
        "destaque": "BEST SELLER"
    },
    {
        "nome": "Grill Elétrico Grande",
        "marca": "BE.SMART",
        "modelo": "BSGR100",
        "especificacoes": "Preto, 127V, Grande Capacidade",
        "preco": 499.00,
        "parcelamento": "1+14x",
        "descricao": "Grelhe à perfeição! Ideal para reuniões com amigos e família. Preparação rápida e resultados deliciosos.",
        "destaque": "ENTRETENIMENTO"
    },
    {
        "nome": "Centrífuga 10kg Cat Comfort",
        "marca": "WANKE",
        "modelo": "CWAT100",
        "especificacoes": "10kg, Branca, 110V",
        "preco": 549.00,
        "parcelamento": "1+14x",
        "descricao": "Seque suas roupas rapidamente com eficiência. Tecnologia Cat Comfort garante cuidado com seus tecidos.",
        "destaque": "QUALIDADE"
    },
    {
        "nome": "Chaleira Elétrica de Vidro com Luz 2L",
        "marca": "POLAK",
        "modelo": "PK7053",
        "especificacoes": "2L, 127V, Vidro Transparente com Iluminação",
        "preco": 179.00,
        "parcelamento": "1+14x",
        "descricao": "Água quentinha em minutos! Design elegante que decora sua cozinha. Perfeita para chás, cafés e bebidas quentes.",
        "destaque": "ELEGÂNCIA"
    },
    {
        "nome": "Caixa de Som Portátil com USB",
        "marca": "ELITE",
        "modelo": "PKR4238",
        "especificacoes": "8W, USB, Portátil",
        "preco": 179.00,
        "parcelamento": "1+14x",
        "descricao": "Leve sua música para qualquer lugar! Som cristalino e bateria durável. Ideal para festas, passeios e viagens.",
        "destaque": "PORTABILIDADE"
    },
    {
        "nome": "Forno Elétrico 48L Plus com Rotisseria",
        "marca": "BEST",
        "modelo": "3411",
        "especificacoes": "48L, Preto/Branco, 127V, Com Rotisseria",
        "preco": 499.00,
        "parcelamento": "1+14x",
        "descricao": "Cozinhe como um profissional! Assados, pizzas e mais. Capacidade grande para toda sua família.",
        "destaque": "VERSATILIDADE"
    },
    {
        "nome": "Fritadeira Elétrica Air Fryer 3,5L",
        "marca": "LIBELL",
        "modelo": "14713",
        "especificacoes": "3,5L, Preta, 127V",
        "preco": 199.00,
        "parcelamento": "1+14x",
        "descricao": "Frituras saudáveis! Com tecnologia Air Fryer, reduz o uso de óleo. Crocante por fora, macio por dentro.",
        "destaque": "SAÚDE"
    },
    {
        "nome": "Sanduicheira Grill 2 em 1",
        "marca": "BE.SMART",
        "modelo": "CM12",
        "especificacoes": "Preta, 127V, Duas Funções",
        "preco": 179.00,
        "parcelamento": "1+14x",
        "descricao": "Sanduíches e grelhados perfeitos! Rápido e prático para o dia a dia. Limpeza fácil e design moderno.",
        "destaque": "PRATICIDADE"
    },
    {
        "nome": "Churrasqueira Tambor Pequena",
        "marca": "Diversos",
        "modelo": "478",
        "especificacoes": "Pequena, Portátil",
        "preco": 178.00,
        "parcelamento": "1+14x",
        "descricao": "Churrasco em qualquer lugar! Perfeita para sítios, churrascos de família e acampamentos. Fácil de transportar.",
        "destaque": "LAZER"
    },
    {
        "nome": "Liquidificador 1,5L 850W",
        "marca": "AGRATTO",
        "modelo": "ALIQ02",
        "especificacoes": "1,5L, 850W, Vermelho, 127V",
        "preco": 899.00,
        "parcelamento": "À vista",
        "descricao": "Bebidas deliciosas em segundos! Potência máxima para sucos, vitaminas e smoothies. Qualidade durável.",
        "destaque": "POTÊNCIA"
    },
    {
        "nome": "TV 32\" Grandes Marcas",
        "marca": "Diversas",
        "modelo": "Selecionadas",
        "especificacoes": "32 polegadas, Full HD, Smart TV",
        "preco": 899.00,
        "parcelamento": "1+14x",
        "descricao": "Imagem nítida e cores vibrantes! Assista seus programas favoritos em qualidade. Marcas conceituadas.",
        "destaque": "ENTRETENIMENTO"
    },
    {
        "nome": "Fogão 4 Bocas Braslar",
        "marca": "BRASLAR",
        "modelo": "6456/322",
        "especificacoes": "4 Bocas, Forno, Aço Inoxidável",
        "preco": 459.00,
        "parcelamento": "1+14x",
        "descricao": "Cozinhe para toda a família! 4 bocas de potência e forno integrado. Acabamento robusto e durável.",
        "destaque": "ESSENCIAL"
    },
    {
        "nome": "Aquecedor de Ar Portátil",
        "marca": "BE.SMART",
        "modelo": "CM8",
        "especificacoes": "127V, Compacto, Eficiente",
        "preco": 89.90,
        "parcelamento": "À vista",
        "descricao": "Aquecimento rápido e confortável! Ideal para salas e quartos. Consumo reduzido de energia.",
        "destaque": "CONFORTO"
    },
    {
        "nome": "Roupeiro 2 Portas Deslizantes",
        "marca": "BRIZ HENN",
        "modelo": "13830",
        "especificacoes": "2 Portas com Corrediça, 2 Gavetas, Espelho",
        "preco": 599.00,
        "parcelamento": "1+14x",
        "descricao": "Organização e elegância no seu quarto! Espaço generoso com portas que deslizam suavemente.",
        "destaque": "ORGANIZAÇÃO"
    },
    {
        "nome": "Cômoda 5 Gavetas",
        "marca": "SIRIUS VILA RICA",
        "modelo": "12108/12109",
        "especificacoes": "5 Gavetas Grandes, Pé Elevado",
        "preco": 599.00,
        "parcelamento": "1+14x",
        "descricao": "Armazene tudo com estilo! Design clássico que combina com qualquer decoração. Gavetas espaçosas.",
        "destaque": "FUNCIONALIDADE"
    },
    {
        "nome": "Estofado 3x2 Lugares",
        "marca": "PARIS ESTOFADOS",
        "modelo": "Paraná",
        "especificacoes": "3x2 Lugares, Tecido Resistente",
        "preco": 1499.00,
        "parcelamento": "1+14x",
        "descricao": "Conforto e sofisticação para sua sala! Assento profundo e encosto aconchegante. Perfeito para relaxar.",
        "destaque": "CONFORTO MÁXIMO"
    },
    {
        "nome": "Roupeiro Fit 2 Portas Deslizantes",
        "marca": "SANTOS ANDIRA",
        "modelo": "13668",
        "especificacoes": "2 Portas Corrediça, Design Moderno",
        "preco": 649.00,
        "parcelamento": "1+14x",
        "descricao": "Moderno e funcional! Perfeito para quartos pequenos. Portas deslizantes economizam espaço.",
        "destaque": "INOVAÇÃO"
    },
    {
        "nome": "Conjunto Box Conjugado Casal",
        "marca": "AC COLCHÕES",
        "modelo": "12749",
        "especificacoes": "Casal 1.38cm, Várias Cores",
        "preco": 599.00,
        "parcelamento": "1+14x",
        "descricao": "Durma melhor e acorde mais feliz! Box bem estruturado para suporte perfeito. Cores sortidas.",
        "destaque": "BOM REPOUSO"
    },
    {
        "nome": "Conjunto Box 138x188cm Mola Ensacada",
        "marca": "APOLO SPUMA",
        "modelo": "14274+14338",
        "especificacoes": "Mola Ensacada Sublime Prime, Premium",
        "preco": 1099.00,
        "parcelamento": "1+14x",
        "descricao": "Tecnologia mola ensacada garante conforto individual. Colchão premium para noites incríveis.",
        "destaque": "PREMIUM"
    },
    {
        "nome": "Kit Cozinha 1.20M Valparaíso",
        "marca": "DACHERI",
        "modelo": "1133/14938",
        "especificacoes": "MDP, Freijo/Baunilha, 1.20M",
        "preco": 599.00,
        "parcelamento": "1+14x",
        "descricao": "Cozinha compacta e bonita! Ideal para apartamentos pequenos. Acabamento elegante em tons naturais.",
        "destaque": "ECONOMIA"
    },
    {
        "nome": "Cozinha Compacta 123cm Sophia",
        "marca": "MÓVEIS SUL",
        "modelo": "14766/14767",
        "especificacoes": "Sem Tampo, 123cm, Compacta",
        "preco": 599.00,
        "parcelamento": "1+14x",
        "descricao": "Solução prática para sua cozinha! Modelo Sophia com design moderno. Sem tampo para flexibilidade.",
        "destaque": "ESPAÇO REDUZIDO"
    },
    {
        "nome": "Cozinha 6 Peças Jade 315cm",
        "marca": "DACHERI",
        "modelo": "14786",
        "especificacoes": "6 Peças, Jade, 315cm, Freijo/Moka",
        "preco": 2299.00,
        "parcelamento": "1+14x",
        "descricao": "Cozinha completa e sofisticada! 6 peças em cores nobres. Ideal para cozinhas maiores e ambientes integrados.",
        "destaque": "COMPLETA"
    },
    {
        "nome": "Armário Multiuso 2 Portas New 900",
        "marca": "ARAMOVEIS",
        "modelo": "13679/14396",
        "especificacoes": "2 Portas, Multiuso, Design New 900",
        "preco": 249.00,
        "parcelamento": "1+14x",
        "descricao": "Organize tudo em um só lugar! Versátil para qualquer cômodo. Qualidade Aramoveis a preço incrível.",
        "destaque": "VERSATILIDADE"
    },
    {
        "nome": "Balcão 1.19M MDP Branco",
        "marca": "DACHERI",
        "modelo": "14287",
        "especificacoes": "1.19M, Branco, Sem Tampo, Pé Elevado",
        "preco": 299.00,
        "parcelamento": "1+14x",
        "descricao": "Bancada moderna para sua cozinha! Pé elevado para limpeza facilitada. Branco que combina com tudo.",
        "destaque": "LIMPEZA FÁCIL"
    },
    {
        "nome": "Pia 1.20M com Válvula",
        "marca": "Marcas Selecionadas",
        "modelo": "Cimentada",
        "especificacoes": "1.20M, Cimentada, Escorredor Integrado",
        "preco": 249.00,
        "parcelamento": "1+14x",
        "descricao": "Pia resistente e durável! Com escorredor integrado. Cimento de qualidade garante longevidade.",
        "destaque": "DURÁVEL"
    },
    {
        "nome": "Jogo de Mesa com 4 Cadeiras Sicília",
        "marca": "VIERO",
        "modelo": "14981",
        "especificacoes": "Mesa Redonda 90x90, Cadeiras Conecta+Dalila",
        "preco": 999.00,
        "parcelamento": "1+14x",
        "descricao": "Reuniões e refeições perfeitas! Mesa redonda aconchegante com 4 cadeiras confortáveis. Estilo clássico.",
        "destaque": "FAMÍLIA"
    },
    {
        "nome": "Painel para TV 1.20M Navi",
        "marca": "BECHARA",
        "modelo": "14957",
        "especificacoes": "1.20M, Navi, Moderno",
        "preco": 289.00,
        "parcelamento": "1+14x",
        "descricao": "Organize sua TV com estilo! Painel moderno que realça seu home theater. Design contemporâneo.",
        "destaque": "HOME THEATER"
    },
    {
        "nome": "Estofado 2.00M Retrátil Reclinável",
        "marca": "MACIEL ESTOFADOS",
        "modelo": "8361",
        "especificacoes": "2.00M, Retrátil, Reclinável, Australia",
        "preco": 1299.00,
        "parcelamento": "1+14x",
        "descricao": "Máximo conforto e funcionalidade! Retrátil para relaxar e reclinável para diferentes posições. Cores sortidas.",
        "destaque": "FUNCIONAL"
    },
]

# Função para adicionar uma linha horizontal
def add_horizontal_line(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'FF6600')
    pBdr.append(bottom)
    pPr.append(pBdr)

# Adicionar cada produto
for idx, produto in enumerate(produtos, 1):
    # Número do anúncio
    numero = doc.add_paragraph()
    numero_run = numero.add_run(f"ANÚNCIO #{idx}")
    numero_run.font.size = Pt(9)
    numero_run.font.bold = True
    numero_run.font.color.rgb = RGBColor(150, 150, 150)

    # Destaque
    destaque = doc.add_paragraph()
    destaque.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    destaque_run = destaque.add_run(f"🌟 {produto['destaque']} 🌟")
    destaque_run.font.size = Pt(11)
    destaque_run.font.bold = True
    destaque_run.font.color.rgb = RGBColor(255, 102, 0)

    # Nome do produto
    nome = doc.add_paragraph()
    nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    nome_run = nome.add_run(produto['nome'])
    nome_run.font.size = Pt(13)
    nome_run.font.bold = True
    nome_run.font.color.rgb = RGBColor(0, 0, 0)

    # Marca e Modelo
    marca = doc.add_paragraph()
    marca.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    marca_run = marca.add_run(f"{produto['marca']} • Modelo: {produto['modelo']}")
    marca_run.font.size = Pt(10)
    marca_run.font.italic = True
    marca_run.font.color.rgb = RGBColor(80, 80, 80)

    # Especificações
    esp = doc.add_paragraph()
    esp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    esp_run = esp.add_run(f"Especificações: {produto['especificacoes']}")
    esp_run.font.size = Pt(9)
    esp_run.font.color.rgb = RGBColor(100, 100, 100)

    # Descrição
    desc = doc.add_paragraph()
    desc.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    desc_run = desc.add_run(f"💬 {produto['descricao']}")
    desc_run.font.size = Pt(10)
    desc_run.font.italic = True
    desc_run.font.color.rgb = RGBColor(60, 60, 60)

    # Espaço
    doc.add_paragraph()

    # Preço (destaque)
    preco_par = doc.add_paragraph()
    preco_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    preco_run = preco_par.add_run(f"R$ {produto['preco']:.2f}")
    preco_run.font.size = Pt(16)
    preco_run.font.bold = True
    preco_run.font.color.rgb = RGBColor(255, 0, 0)

    # Forma de pagamento
    pag = doc.add_paragraph()
    pag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    pag_run = pag.add_run(f"Parcelado em: {produto['parcelamento']}")
    pag_run.font.size = Pt(10)
    pag_run.font.bold = True
    pag_run.font.color.rgb = RGBColor(0, 102, 204)

    # Contato
    contato = doc.add_paragraph()
    contato.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contato_run = contato.add_run("📞 (42) 99989-0217 • @lojasmdloficial")
    contato_run.font.size = Pt(9)
    contato_run.font.color.rgb = RGBColor(100, 100, 100)

    # Linha separadora
    separador = doc.add_paragraph()
    add_horizontal_line(separador)

    # Espaço entre anúncios
    doc.add_paragraph()

# Adicionar footer
footer_par = doc.add_paragraph()
footer_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
footer_run = footer_par.add_run("MDL - Móveis do Lar\nOferta válida para Maio de 2026 ou enquanto durar o estoque!\nCrédito próprio em até 15x • Frete disponível • Quantidade limitada")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(100, 100, 100)
footer_run.font.italic = True

# Salvar documento
doc.save('Anuncios_Individuais_Produtos_MDL.docx')
print("✓ Documento criado com sucesso: Anuncios_Individuais_Produtos_MDL.docx")
print(f"✓ Total de {len(produtos)} anúncios gerados!")
